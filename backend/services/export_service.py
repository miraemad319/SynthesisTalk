from fpdf import FPDF
from docx import Document
from typing import Dict, Any
import unicodedata
import base64
import io
from PIL import Image

def _clean_text_for_pdf(text: str) -> str:
    """Clean text to remove characters that can't be encoded in Latin-1."""
    if not text:
        return ""
    
    # Replace common Unicode characters with ASCII equivalents
    replacements = {
        '↔': '<->',
        '→': '->',
        '←': '<-',
        '…': '...',
        '"': '"',
        '"': '"',
        ''': "'",
        ''': "'",
        '–': '-',
        '—': '--'
    }
    
    for unicode_char, ascii_char in replacements.items():
        text = text.replace(unicode_char, ascii_char)
    
    # Remove any remaining non-Latin-1 characters
    try:
        text.encode('latin-1')
        return text
    except UnicodeEncodeError:
        # If there are still problematic characters, normalize and filter
        normalized = unicodedata.normalize('NFKD', text)
        ascii_text = ''.join(c for c in normalized if ord(c) < 256)
        return ascii_text

def export_findings(insights_report: Dict[str, Any], file_path: str, format: str = "pdf") -> None:
    """
    Export findings to a standard document format (PDF or Word).

    Args:
        insights_report (Dict[str, Any]): The insights report to export.
        file_path (str): The file path to save the document.
        format (str): The format to export ("pdf" or "word").
    """
    if format.lower() == "pdf":
        _export_to_pdf(insights_report, file_path)
    elif format.lower() == "word":
        _export_to_word(insights_report, file_path)
    else:
        raise ValueError("Unsupported format. Please choose 'pdf' or 'word'.")

def _export_to_pdf(insights_report: Dict[str, Any], file_path: str) -> None:
    """Export findings to a PDF file."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    # Add title
    pdf.set_font("Arial", style="B", size=16)
    pdf.cell(200, 10, txt="Insights Report", ln=True, align="C")
    pdf.ln(10)

    # Add summary
    pdf.set_font("Arial", size=12)
    summary_text = _clean_text_for_pdf(insights_report.get('summary', ''))
    pdf.multi_cell(0, 10, txt=f"Summary:\n{summary_text}")
    pdf.ln(5)

    # Add key patterns
    pdf.set_font("Arial", style="B", size=14)
    pdf.cell(0, 10, txt="Key Patterns:", ln=True)
    pdf.set_font("Arial", size=12)
    for pattern in insights_report.get("key_patterns", []):
        pattern_text = _clean_text_for_pdf(f"- {pattern['term']} (Frequency: {pattern['frequency']})")
        pdf.cell(0, 10, txt=pattern_text, ln=True)
    pdf.ln(5)

    # Add trends
    pdf.set_font("Arial", style="B", size=14)
    pdf.cell(0, 10, txt="Trends:", ln=True)
    pdf.set_font("Arial", size=12)
    for trend in insights_report.get("trends", []):
        trend_text = _clean_text_for_pdf(f"- {trend['type'].capitalize()} ({trend['indicator']})")
        pdf.cell(0, 10, txt=trend_text, ln=True)
    pdf.ln(5)

    # Add relationships
    pdf.set_font("Arial", style="B", size=14)
    pdf.cell(0, 10, txt="Relationships:", ln=True)
    pdf.set_font("Arial", size=12)
    for relationship in insights_report.get("relationships", []):
        # Use ASCII arrow instead of Unicode
        relationship_text = _clean_text_for_pdf(f"- {relationship['term1']} <-> {relationship['term2']} (Strength: {relationship['strength']})")
        pdf.cell(0, 10, txt=relationship_text, ln=True)
    pdf.ln(5)
    
    # Add visualizations
    visualizations = insights_report.get("visualizations", [])
    if visualizations:
        pdf.set_font("Arial", style="B", size=14)
        pdf.cell(0, 10, txt="Visualizations:", ln=True)
        pdf.ln(5)
        
        for i, viz in enumerate(visualizations):
            if 'data' in viz and viz['data']:
                try:
                    # Decode base64 image
                    img_data = base64.b64decode(viz['data'])
                    img_file = f"temp_viz_{i}.png"
                    
                    # Save temporarily
                    with open(img_file, 'wb') as f:
                        f.write(img_data)
                    
                    # Add title for the visualization
                    pdf.set_font("Arial", style="I", size=12)
                    pdf.cell(0, 10, txt=viz.get('title', f"Visualization {i+1}"), ln=True)
                    
                    # Add the image
                    try:
                        pdf.image(img_file, x=10, w=180)
                        pdf.ln(5)
                        
                        # Add description if available
                        if 'description' in viz:
                            pdf.set_font("Arial", size=10)
                            description_text = _clean_text_for_pdf(viz['description'])
                            pdf.multi_cell(0, 5, txt=description_text)
                            pdf.ln(5)
                    except Exception as e:
                        pdf.cell(0, 10, txt=f"[Error displaying visualization: {str(e)}]", ln=True)
                    
                    # Clean up temp file
                    import os
                    if os.path.exists(img_file):
                        os.remove(img_file)
                        
                except Exception as e:
                    pdf.cell(0, 10, txt=f"[Error processing visualization: {str(e)}]", ln=True)

    # Save PDF
    pdf.output(file_path)

def _export_to_word(insights_report: Dict[str, Any], file_path: str) -> None:
    """Export findings to a Word document."""
    doc = Document()

    # Add title
    doc.add_heading("Insights Report", level=1)

    # Add summary
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(insights_report.get("summary", ""))

    # Add key patterns
    doc.add_heading("Key Patterns", level=2)
    for pattern in insights_report.get("key_patterns", []):
        doc.add_paragraph(f"- {pattern['term']} (Frequency: {pattern['frequency']})")

    # Add trends
    doc.add_heading("Trends", level=2)
    for trend in insights_report.get("trends", []):
        doc.add_paragraph(f"- {trend['type'].capitalize()} ({trend['indicator']})")

    # Add relationships
    doc.add_heading("Relationships", level=2)
    for relationship in insights_report.get("relationships", []):
        # Use ASCII arrow for consistency
        doc.add_paragraph(f"- {relationship['term1']} <-> {relationship['term2']} (Strength: {relationship['strength']})")
    
    # Add visualizations
    visualizations = insights_report.get("visualizations", [])
    if visualizations:
        doc.add_heading("Visualizations", level=2)
        
        for i, viz in enumerate(visualizations):
            if 'data' in viz and viz['data']:
                try:
                    # Decode base64 image
                    img_data = base64.b64decode(viz['data'])
                    img_file = f"temp_viz_{i}.png"
                    
                    # Save temporarily
                    with open(img_file, 'wb') as f:
                        f.write(img_data)
                    
                    # Add title for the visualization
                    doc.add_heading(viz.get('title', f"Visualization {i+1}"), level=3)
                    
                    # Add the image to the document
                    try:
                        doc.add_picture(img_file, width=6000000)  # Width in EMU (English Metric Units)
                        
                        # Add description if available
                        if 'description' in viz:
                            doc.add_paragraph(viz['description'], style='Italic')
                    except Exception as e:
                        doc.add_paragraph(f"[Error displaying visualization: {str(e)}]")
                    
                    # Clean up temp file
                    import os
                    if os.path.exists(img_file):
                        os.remove(img_file)
                        
                except Exception as e:
                    doc.add_paragraph(f"[Error processing visualization: {str(e)}]")

    # Save Word document
    doc.save(file_path)